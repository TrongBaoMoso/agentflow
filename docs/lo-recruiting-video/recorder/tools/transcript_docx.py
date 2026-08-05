"""Renders the payload from build-transcript-doc.mjs into a landscape .docx.

Called by that script; not meant to be run on its own. Kept in Python because python-docx is the only
reliable way to put an image *inside a table cell* in a file Drive will convert in place.

Layout contract: one row per scene, image left at a fixed 5.7in (so every screenshot is the same
size and the eye can compare boards across acts), narration right. Rows are marked cantSplit so a
frame is never on one page with its narration on the next.
"""
import json
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0xE8, 0x57, 0x0E)   # LF orange — act headings only
VI_INK = RGBColor(0x8A, 0x3B, 0x0C)   # darker orange so Vietnamese reads as a second voice
META = RGBColor(0x77, 0x77, 0x77)
BODY = RGBColor(0x22, 0x22, 0x22)
NOTE_INK = RGBColor(0x8A, 0x6D, 0x1F)

IMG_W = Inches(5.7)
COL_IMG = Inches(5.9)
COL_TXT = Inches(3.9)


def borders(table, color="D9D9D9", sz=4):
    tbl_pr = table._tbl.tblPr
    el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        el.append(e)
    tbl_pr.append(el)


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def para(container, text="", *, size=10, color=BODY, font="Arial", bold=False,
         italic=False, before=0, after=4, spacing=1.15, align=None):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = spacing
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = font
        r.bold = bold
        r.italic = italic
    return p


def build(payload):
    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    film = payload["film"]

    # ---- cover -------------------------------------------------------------------------------
    para(doc, "LO RECRUITING — HỆ THỐNG ĐANG CHẠY", size=9, color=ACCENT, bold=True, after=2)
    para(doc, "Bản text của video walkthrough theo role", size=26, color=BODY,
         font="Georgia", bold=True, after=6, spacing=1.0)
    para(doc,
         f"{film['scenes']} cảnh · {film['acts']} act · {film['length']} phim · quay trên PRODUCTION "
         f"({film['file']})",
         size=10, color=META, after=14)

    para(doc, "Đây là gì", size=11, color=BODY, font="Georgia", bold=True, after=3)
    para(doc,
         "Mỗi dòng dưới đây là một cảnh trong phim: bên trái là màn hình thật tại thời điểm đó, bên "
         "phải là đúng lời thuyết minh của cảnh đó — tiếng Anh (giọng đọc trong phim) ở trên, tiếng Việt "
         "(phụ đề) ở dưới. Không có cảnh nào được dựng, không có dữ liệu nào được bịa: record thật, "
         "người thật, số thật, đo tại thời điểm quay 05/08/2026.",
         after=6)
    para(doc,
         "Mốc thời gian bên cạnh mã cảnh (ví dụ 10:26) là giây mà cảnh đó bắt đầu trong phim — mở video "
         "tới đúng giây đó là thấy nguyên cảnh, kèm âm thanh.",
         size=9, color=META, after=14)

    para(doc, "Bảy vai, cùng một hệ thống", size=11, color=BODY, font="Georgia", bold=True, after=4)
    cast = doc.add_table(rows=0, cols=3)
    borders(cast)
    for a in payload["acts"]:
        row = cast.add_row()
        no_split(row)
        widths = (Inches(2.6), Inches(2.8), Inches(4.4))
        for cell, w in zip(row.cells, widths):
            cell.width = w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        para(row.cells[0], a["title"], size=10, bold=True, after=0)
        para(row.cells[1], a["who"], size=10, color=META, after=0)
        para(row.cells[2], a["gloss"], size=10, color=BODY, after=0)
    cast.autofit = False

    para(doc,
         "Con số quyền là số công tắc được tick trên đúng cây 82 quyền của hệ thống, đọc trực tiếp từ "
         "trang Associates của từng người. Cùng một cây, sáu con số khác nhau: quyền ở đây được cấp theo "
         "từng người, không suy ra được từ vai.",
         size=9, color=META, before=6, after=0)

    # ---- one table per act -------------------------------------------------------------------
    by_act = {}
    for s in payload["scenes"]:
        by_act.setdefault(s["act"], []).append(s)

    for a in payload["acts"]:
        scenes = by_act.get(a["act"], [])
        if not scenes:
            continue
        doc.add_page_break()
        para(doc, a["title"].upper(), size=15, color=ACCENT, font="Georgia", bold=True, after=1)
        para(doc, f"{a['who']} — {a['gloss']}", size=10, color=META, after=8)

        table = doc.add_table(rows=0, cols=2)
        borders(table)
        table.autofit = False
        for s in scenes:
            row = table.add_row()
            no_split(row)
            left, right = row.cells
            left.width, right.width = COL_IMG, COL_TXT
            for cell in (left, right):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)

            p = para(left, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run().add_picture(s["file"], width=IMG_W)

            head = f"{s['id']}  ·  {s['stamp']}  ·  {s['dur']}"
            if s.get("shot"):
                head += f"  ·  {s['shot']}"
            para(right, head, size=8, color=META, bold=True, after=5)
            para(right, s["en"], size=10, font="Georgia", color=BODY, after=6)
            para(right, s["vi"], size=10, color=VI_INK, after=0)
            if s["note"]:
                para(right, s["note"], size=8, color=NOTE_INK, italic=True, before=6, after=0)

    doc.save(payload["outDocx"])
    print(f"   docx    {payload['outDocx']}")


if __name__ == "__main__":
    with open(sys.argv[1], encoding="utf-8") as fh:
        build(json.load(fh))
